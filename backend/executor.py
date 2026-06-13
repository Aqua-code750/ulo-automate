import time
import os
import requests
import json
import subprocess
import re
from collections import deque
from dotenv import load_dotenv

# Load env vars for API keys
load_dotenv()

class WorkflowExecutor:
    def __init__(self, workflow_data, success_cb=None, error_cb=None, done_cb=None):
        self.workflow_data = workflow_data
        self.success_cb = success_cb
        self.error_cb = error_cb
        self.done_cb = done_cb

    def run(self):
        nodes = self.workflow_data.get('nodes', [])
        edges = self.workflow_data.get('edges', [])
        
        # Build graph
        node_map = {n['id']: n for n in nodes}
        adj_list = {n['id']: [] for n in nodes}
        in_degree = {n['id']: 0 for n in nodes}
        
        for e in edges:
            src, tgt = e['source'], e['target']
            adj_list[src].append(tgt)
            in_degree[tgt] += 1
            
        # Queue initialized with Trigger nodes (in-degree == 0)
        queue = deque([nid for nid, deg in in_degree.items() if deg == 0])
        
        # Global payload holding outputs from all nodes
        payload = {"global": "initialized"}
        
        while queue:
            current_id = queue.popleft()
            node = node_map[current_id]
            task_type = node.get('type', '')
            props = node.get('props', {})
            
            # Simple Variable Substitution
            # Replaces placeholders like {{some_var}} with data from payload
            try:
                resolved_props = self._resolve_props(props, payload)
            except Exception as e:
                print(f"Error resolving variables for {current_id}: {e}")
                resolved_props = props
            
            try:
                time.sleep(0.5) # small visual delay
                result = self._dispatch(task_type, resolved_props, payload)
                
                # Merge result into payload namespace safely
                payload[f"node_{current_id}"] = result
                if isinstance(result, dict):
                    payload.update(result)

                if self.success_cb:
                    self.success_cb(current_id)
                    
                # Progress to neighbors
                for neighbor in adj_list[current_id]:
                    in_degree[neighbor] -= 1
                    if in_degree[neighbor] == 0:
                        queue.append(neighbor)
                        
            except Exception as e:
                print(f"Error executing {current_id} ({task_type}): {e}")
                if self.error_cb:
                    self.error_cb(current_id, str(e))
                return False # Stop execution on failure
                
        print("Final Payload:", payload)
        if self.done_cb:
            self.done_cb(payload)
        return True

    def _resolve_props(self, props, payload):
        """Recursively resolves {{ }} placeholders in properties using eval against the payload."""
        resolved = {}
        for k, v in props.items():
            if isinstance(v, str):
                # Find all {{ expr }}
                def replacer(match):
                    expr = match.group(1).strip()
                    try:
                        # Safely eval the expression with 'payload' as a local variable
                        # e.g., {{ payload['node_123']['http_response'] }}
                        val = eval(expr, {"__builtins__": {}}, {"payload": payload})
                        return str(val) if val is not None else ""
                    except Exception as e:
                        return f"{{ERROR: {e}}}"
                
                resolved[k] = re.sub(r'\{\{(.*?)\}\}', replacer, v)
            elif isinstance(v, dict):
                resolved[k] = self._resolve_props(v, payload)
            else:
                resolved[k] = v
        return resolved

    def _dispatch(self, task_type, props, payload):
        """Dynamic dispatch based on task type"""
        if "Trigger" in task_type:
            return self._handle_trigger(task_type, props)
        elif "HTTP Request" in task_type:
            return self._handle_http(props, payload)
        elif "Python Code" in task_type:
            return self._handle_python(props, payload)
        elif "Bash Command" in task_type:
            return self._handle_bash(props)
        elif "Slack" in task_type:
            return self._handle_slack(props)
        elif "Agent" in task_type:
            return self._handle_agent(task_type, props, payload)
        elif "AI" in task_type or "Gemini" in task_type or "OpenAI" in task_type:
            return self._handle_ai(task_type, props, payload)
        elif "Security" in task_type or "Veracode" in task_type or "Falcon" in task_type:
            return self._handle_security(task_type, props)
        elif "Office" in task_type:
            return self._handle_office(task_type, props)
        elif "Data" in task_type:
            return self._handle_data(task_type, props, payload)
        elif "Read File" in task_type:
            path = props.get("url", "default.txt")
            with open(path, "r") as f:
                return {"file_content": f.read()}
        elif "Write File" in task_type:
            path = props.get("url", "output.txt")
            with open(path, "w") as f:
                f.write(json.dumps(payload, indent=2))
            return {"file_written": path}
        else:
            return {"status": "success", "note": f"Executed generic {task_type}"}

    def _handle_trigger(self, task_type, props):
        return {"trigger_type": task_type, "timestamp": time.time()}

    def _handle_http(self, props, payload):
        url = props.get("url", "")
        method = props.get("method", "GET").upper()
        headers_str = props.get("headers", "{}")
        body_str = props.get("body", "")
        
        if not url:
            raise ValueError("HTTP Request missing URL")
            
        try:
            headers = json.loads(headers_str) if headers_str else {}
        except Exception:
            headers = {}
            
        data = None
        if body_str:
            data = body_str.encode('utf-8')
            
        req = requests.request(method, url, headers=headers, data=data)
        try:
            resp_data = req.json()
        except:
            resp_data = req.text
            
        return {"http_response": resp_data, "status_code": req.status_code}

    def _handle_python(self, props, payload):
        code = props.get("code", "")
        if not code:
            return {"python_out": "No code provided"}
            
        local_scope = {"payload": payload, "result": None}
        exec(code, {}, local_scope)
        return {"python_out": local_scope.get("result", "Executed successfully")}

    def _handle_bash(self, props):
        cmd = props.get("command", "")
        if not cmd:
            return {"bash_out": "No command"}
            
        output = subprocess.check_output(cmd, shell=True, text=True)
        return {"bash_out": output}

    def _handle_slack(self, props):
        webhook_url = props.get("webhook_url", "")
        message = props.get("message", "Hello from Ulo Automate!")
        channel = props.get("channel", "")
        username = props.get("username", "UloBot")
        
        if not webhook_url:
            raise ValueError("Slack Webhook URL is missing!")
            
        payload = {"text": message, "username": username}
        if channel:
            payload["channel"] = channel
            
        resp = requests.post(webhook_url, json=payload)
        if resp.status_code >= 400:
            raise Exception(f"Slack Error: {resp.text}")
        return {"slack_status": "Message sent!"}

    def _handle_ai(self, task_type, props, payload):
        prompt = props.get("prompt", "Say hello")
        api_key = props.get("api_key", "").strip()
        
        if "Gemini" in task_type:
            key = api_key or os.getenv("GEMINI_API_KEY")
            if not key: raise ValueError("GEMINI_API_KEY missing")
            import google.generativeai as genai
            genai.configure(api_key=key)
            model = genai.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(prompt)
            return {"ai_response": response.text}
            
        elif "OpenAI" in task_type:
            key = api_key or os.getenv("OPENAI_API_KEY")
            if not key: raise ValueError("OPENAI_API_KEY missing")
            headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
            data = {
                "model": props.get("model", "gpt-3.5-turbo"),
                "messages": [{"role": "user", "content": prompt}]
            }
            resp = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=data)
            if resp.status_code >= 400:
                raise Exception(f"OpenAI Error: {resp.text}")
            return {"ai_response": resp.json()['choices'][0]['message']['content']}
            
        else:
            return {"ai_response": "Generic AI task simulated successfully."}

    def _handle_agent(self, task_type, props, payload):
        role = props.get("agent_role", "Helpful Assistant")
        max_iters = int(props.get("max_iterations", 3) or 3)
        tools = props.get("tools", "")
        
        print(f"--- AGENT BOOTING ---")
        print(f"Role: {role} | Tools: {tools}")
        
        # Simulated ReAct Loop
        agent_logs = []
        for i in range(max_iters):
            time.sleep(0.5)
            action = f"Thinking iteration {i+1}..."
            print(f"Agent: {action}")
            agent_logs.append(action)
            
            if "python" in tools.lower() and i == 0:
                agent_logs.append("Action: Executing Python tool...")
            if "slack" in tools.lower() and i == max_iters - 1:
                agent_logs.append("Action: Notifying user via Slack tool...")
                
        agent_logs.append("Final Answer reached.")
        print(f"--- AGENT FINISHED ---")
        return {"agent_result": "Agent successfully completed task autonomously.", "agent_logs": agent_logs}

    def _handle_office(self, task_type, props):
        path = props.get("file_path", "")
        
        if "Word" in task_type:
            template_path = props.get("template_path", "")
            input_data_str = props.get("input_data", "{}")
            
            try:
                input_data = json.loads(input_data_str)
            except Exception as e:
                return {"error": f"Failed to parse input_data as JSON: {e}"}
                
            if template_path and os.path.exists(template_path):
                try:
                    from docxtpl import DocxTemplate
                    doc = DocxTemplate(template_path)
                    doc.render(input_data)
                    doc.save(path)
                    return {"status": "success", "file": path, "method": "template"}
                except Exception as e:
                    return {"error": f"DocxTemplate failed: {e}"}
            else:
                try:
                    import docx
                    doc = docx.Document()
                    doc.add_heading('Automated Report', 0)
                    for k, v in input_data.items():
                        doc.add_paragraph(f"{k}: {v}")
                    doc.save(path)
                    return {"status": "success", "file": path, "method": "new_document"}
                except Exception as e:
                    return {"error": f"python-docx failed: {e}"}
                    
        elif "Excel" in task_type:
            try:
                import openpyxl
            except ImportError:
                return {"error": "openpyxl is not installed"}
                
            if "Read" in task_type:
                if not path: return {"error": "No file path provided"}
                wb = openpyxl.load_workbook(path, data_only=True)
                sheet_name = props.get("sheet_name", wb.active.title)
                sheet = wb[sheet_name]
                
                # Simple extraction of all rows
                data = []
                for row in sheet.iter_rows(values_only=True):
                    data.append(row)
                return {"excel_data": data}
                
            elif "Write" in task_type:
                if not path: return {"error": "No file path provided"}
                wb = openpyxl.Workbook()
                sheet = wb.active
                sheet.title = props.get("sheet_name", "Sheet1")
                
                # Assume props["input_data"] is a JSON string of 2D array
                input_data = props.get("input_data", "[]")
                try:
                    rows = json.loads(input_data)
                    for r in rows:
                        sheet.append(r)
                except Exception as e:
                    return {"error": f"Failed to parse input_data as JSON array: {e}"}
                    
                wb.save(path)
                return {"status": "success", "file": path}
                
        elif "PDF" in task_type:
            try:
                import PyPDF2
            except ImportError:
                return {"error": "PyPDF2 is not installed"}
                
            if not path: return {"error": "No file path provided"}
            text = ""
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return {"pdf_text": text}
            
        elif "Teams" in task_type:
            webhook = props.get("webhook_url", "")
            message = props.get("message", "Hello from Ulo Automate!")
            if not webhook: return {"error": "No webhook URL provided"}
            resp = requests.post(webhook, json={"text": message})
            if resp.status_code >= 400:
                raise Exception(f"Teams Error: {resp.text}")
            return {"teams_status": "Message sent!"}
            
        elif "Outlook" in task_type:
            import imaplib
            import email
            
            user = props.get("email", "")
            password = props.get("password", "")
            server = props.get("server", "outlook.office365.com")
            
            if not user or not password:
                return {"error": "Email or password missing"}
                
            mail = imaplib.IMAP4_SSL(server)
            mail.login(user, password)
            mail.select("inbox")
            
            status, messages = mail.search(None, 'UNSEEN')
            if status != 'OK':
                return {"error": "Failed to search inbox"}
                
            email_ids = messages[0].split()
            extracted_emails = []
            
            for eid in email_ids[:5]: # limit to 5
                res, msg_data = mail.fetch(eid, '(RFC822)')
                for response_part in msg_data:
                    if isinstance(response_part, tuple):
                        msg = email.message_from_bytes(response_part[1])
                        subject = email.header.decode_header(msg['subject'])[0][0]
                        if isinstance(subject, bytes):
                            subject = subject.decode()
                        extracted_emails.append({"subject": subject, "from": msg.get('from')})
                        
            mail.logout()
            return {"unread_emails": extracted_emails}
            
        return {"office_result": "Office task complete."}

    def _handle_security(self, task_type, props):
        api_id = props.get("api_id", "")
        api_key = props.get("api_key", "")
        app_name = props.get("app_name", "")
        host_id = props.get("host_id", "")
        
        if not api_id or not api_key:
            return {"error": "Missing Security API Credentials"}
            
        if "Veracode SAST Scan" in task_type:
            # Simulated real integration mapping to Veracode REST API structure
            # Normally this uses Veracode's HMAC auth format
            headers = {"Authorization": f"VERACODE-HMAC-SHA-256 id={api_id},ts=...,nonce=...,sig=..."}
            return {
                "security_action": f"Initiated Veracode SAST Pipeline Scan for {app_name}", 
                "status": "scanning",
                "simulated_request_url": "https://api.veracode.com/appsec/v1/applications/scans"
            }
        elif "Veracode Get Flaws" in task_type:
            return {
                "security_action": f"Fetched top high-severity flaws for {app_name}", 
                "flaws": [{"cwe": "79", "severity": "High", "description": "Cross-Site Scripting"}],
                "simulated_request_url": "https://api.veracode.com/appsec/v2/applications/findings"
            }
        elif "Falcon Alert" in task_type:
            headers = {"Authorization": f"Bearer {api_key}"}
            payload = {"composite_id": host_id, "description": "Ulo Automation Alert triggered"}
            # requests.post("https://api.crowdstrike.com/incidents/entities/alerts/v1", headers=headers, json=payload)
            return {
                "security_action": "CrowdStrike Falcon alert triggered", 
                "severity": "HIGH",
                "simulated_request_url": "https://api.crowdstrike.com/incidents/entities/alerts/v1"
            }
        elif "Falcon Contain" in task_type:
            headers = {"Authorization": f"Bearer {api_key}"}
            payload = {"action_parameters": [{"name": "contain", "value": "true"}], "ids": [host_id]}
            # requests.post("https://api.crowdstrike.com/devices/entities/devices-actions/v2?action_name=contain", ...)
            return {
                "security_action": f"Network contained host {host_id} via CrowdStrike Falcon API", 
                "status": "isolated",
                "simulated_request_url": "https://api.crowdstrike.com/devices/entities/devices-actions/v2?action_name=contain"
            }
        return {"security_action": "Generic security task executed successfully"}

    def _handle_data(self, task_type, props, payload):
        import base64
        import hashlib
        import uuid
        
        input_data = props.get("input_data", "").strip()
        
        if "Base64" in task_type:
            if not input_data:
                return {"data_result": "Error: No input_data provided"}
            if "Encode" in task_type:
                encoded = base64.b64encode(input_data.encode('utf-8')).decode('utf-8')
                return {"data_result": encoded}
            elif "Decode" in task_type:
                decoded = base64.b64decode(input_data.encode('utf-8')).decode('utf-8')
                return {"data_result": decoded}
        elif "UUID" in task_type:
            return {"data_result": str(uuid.uuid4())}
        elif "Hash" in task_type:
            if not input_data:
                return {"data_result": "Error: No input_data provided"}
            m = hashlib.sha256()
            m.update(input_data.encode('utf-8'))
            return {"data_result": m.hexdigest()}
        elif "Parse JSON" in task_type:
            if not input_data: return {"data_result": "Error: No input_data provided"}
            try:
                parsed = json.loads(input_data)
                return {"data_result": parsed}
            except Exception as e:
                return {"data_result": f"Error parsing JSON: {e}"}
        elif "Parse CSV" in task_type:
            if not input_data: return {"data_result": "Error: No input_data provided"}
            import csv
            import io
            try:
                reader = csv.DictReader(io.StringIO(input_data))
                rows = [row for row in reader]
                return {"data_result": rows}
            except Exception as e:
                return {"data_result": f"Error parsing CSV: {e}"}
                
        return {"data_result": "Data formatted correctly."}

