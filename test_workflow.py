import os
import sys

# Ensure we can import from backend
sys.path.append(os.path.abspath("backend"))

from docx import Document

# 1. Create a simple Word template with docxtpl/Jinja2 syntax
doc = Document()
doc.add_heading('User Logs Report', 0)
doc.add_paragraph('Here is the simplified list of users from the logs:')

# Add a Jinja2 loop for docxtpl
doc.add_paragraph('{% for user in users %}')
doc.add_paragraph('• Name: {{ user.name }} | Role: {{ user.role }} | Last Login: {{ user.login }}')
doc.add_paragraph('{% endfor %}')

template_path = os.path.abspath("user_template.docx")
doc.save(template_path)
print(f"Created template at: {template_path}")

# 2. Build the workflow payload simulating Ulo Automate nodes
workflow_nodes = [
    {
        "id": "node_1",
        "type": "Trigger: Manual",
        "props": {
            "mock_data": "start"
        }
    },
    {
        "id": "node_2",
        "type": "Data: Parse JSON",
        "props": {
            "json_string": '{"users": [{"name": "Alice", "role": "Admin", "login": "2026-06-13"}, {"name": "Bob", "role": "User", "login": "2026-06-12"}, {"name": "Charlie", "role": "Guest", "login": "2026-06-10"}]}'
        }
    },
    {
        "id": "node_3",
        "type": "Office: Write to Excel",
        "props": {
            "file_path": os.path.abspath("users_log.xlsx"),
            "sheet_name": "Logs",
            "input_data": '[["Name", "Role", "Last Login"], ["Alice", "Admin", "2026-06-13"], ["Bob", "User", "2026-06-12"], ["Charlie", "Guest", "2026-06-10"]]'
        }
    },
    {
        "id": "node_4",
        "type": "Office: Generate Word Document",
        "props": {
            "file_path": os.path.abspath("final_user_report.docx"),
            "template_path": template_path,
            # We pass the parsed JSON object directly using our variable substitution!
            "input_data": '{{ payload["node_2"] }}'
        }
    }
]

print("Running Workflow through Ulo Automate Executor...")
from backend.executor import WorkflowExecutor

executor = WorkflowExecutor()
results = executor.execute(workflow_nodes)

print("\nWorkflow Execution Results:")
for node_id, result in results.items():
    print(f"{node_id}: {result}")

print("\nFiles Created:")
print("-", os.path.abspath("users_log.xlsx"))
print("-", os.path.abspath("final_user_report.docx"))
